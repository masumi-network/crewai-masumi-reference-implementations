from crewai import Agent, Crew, Task
from textwrap import dedent
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import os
import uuid
from datetime import datetime
from enum import Enum
from typing import Optional, Dict, Any
from pydantic import BaseModel, Field

class ContractType(str, Enum):
    NDA = "nda"
    FREELANCE = "freelance"
    EMPLOYMENT = "employment"

class ContractDetails(BaseModel):
    contract_type: ContractType
    company_name: str
    company_address: str
    
    # Party Information
    party_name: str
    party_address: str
    party_email: str
    
    # Contract Specific Details
    start_date: str
    
    # Optional fields based on contract type
    end_date: Optional[str] = None
    salary: Optional[float] = None  # For employment
    hourly_rate: Optional[float] = None  # For freelance
    project_scope: Optional[str] = None  # For freelance
    confidentiality_period: Optional[str] = None  # For NDA
    jurisdiction: str = "California, USA"  # Default jurisdiction
    
    # Additional terms
    additional_terms: Optional[str] = None

class ContractCreationCrew:
    def __init__(self, verbose=True):
        self.verbose = verbose
        self.crew = self.create_crew()
        # Create a directory for storing PDFs if it doesn't exist
        self.pdf_dir = "generated_contracts"
        if not os.path.exists(self.pdf_dir):
            os.makedirs(self.pdf_dir)

    def generate_pdf(self, content, filename=None):
        """Generate PDF from contract content"""
        if filename is None:
            filename = f"contract_{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        filepath = os.path.join(self.pdf_dir, filename)
        
        # First create a Word document for better formatting
        doc = Document()
        
        # Add title
        doc.add_heading('Contract Agreement', 0)
        
        # Split content into sections and add to document
        sections = content.split('\n\n')
        for section in sections:
            if section.strip():
                doc.add_paragraph(section.strip())
        
        # Save the Word document temporarily
        temp_docx = os.path.join(self.pdf_dir, "temp.docx")
        doc.save(temp_docx)
        
        # Convert to PDF
        c = canvas.Canvas(filepath, pagesize=letter)
        c.setFont("Helvetica", 12)
        
        # Add content to PDF
        y = 750  # Starting y position
        for section in sections:
            if section.strip():
                # Split long text into lines that fit on the page
                words = section.strip().split()
                line = []
                for word in words:
                    line.append(word)
                    if len(' '.join(line)) > 65:  # Adjust number based on page width
                        c.drawString(72, y, ' '.join(line[:-1]))
                        line = [line[-1]]
                        y -= 20  # Line spacing
                        
                        # Check if we need a new page
                        if y < 50:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y = 750
                
                if line:
                    c.drawString(72, y, ' '.join(line))
                y -= 30  # Paragraph spacing
                
                # Check if we need a new page
                if y < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = 750
        
        c.save()
        
        # Clean up temporary file
        os.remove(temp_docx)
        
        # Return the relative path to the PDF
        return filepath

    def create_crew(self):
        # Legal Expert Agent with enhanced specialization
        legal_expert = Agent(
            role='Legal Compliance Expert',
            goal='Ensure contract compliance with employment, freelance, and NDA regulations',
            backstory=dedent("""
                Senior legal professional with 15+ years of experience in employment law, 
                contractor agreements, and confidentiality contracts. Expert in California 
                labor laws, intellectual property rights, and non-disclosure agreements. 
                Specialized in ensuring contracts meet both federal and state requirements.
            """),
            verbose=self.verbose,
            allow_delegation=False
        )

        # Contract Specialist with type-specific expertise
        contract_specialist = Agent(
            role='Contract Creation Specialist',
            goal='Create specialized contracts based on type and requirements',
            backstory=dedent("""
                Expert contract writer with specific expertise in NDAs, employment contracts, 
                and freelance agreements. Skilled at incorporating industry-specific terms 
                and protecting both parties' interests. Experienced in technology sector 
                contracts and California employment law requirements.
            """),
            verbose=self.verbose,
            allow_delegation=True
        )

        # Risk Assessment Agent
        risk_analyst = Agent(
            role='Risk Assessment Specialist',
            goal='Analyze contract-specific risks based on type and jurisdiction',
            backstory=dedent("""
                Risk assessment expert specialized in employment law, contractor relationships, 
                and intellectual property protection. Experienced in identifying potential 
                legal vulnerabilities in various contract types and suggesting appropriate 
                safeguards. Expert in California employment risk mitigation.
            """),
            verbose=self.verbose,
            allow_delegation=True
        )

        # Final Reviewer Agent
        final_reviewer = Agent(
            role='Contract Review Specialist',
            goal='Ensure contract type-specific compliance and completeness',
            backstory=dedent("""
                Senior contract reviewer with expertise in all three contract types (NDA, 
                Employment, Freelance). Specialized in ensuring each contract type meets 
                its specific legal requirements and industry standards. Expert in final 
                validation of contract terms and conditions.
            """),
            verbose=self.verbose,
            allow_delegation=False
        )

        # Document Formatter Agent
        document_formatter = Agent(
            role='Document Formatting Specialist',
            goal='Format contracts according to type-specific standards',
            backstory=dedent("""
                Expert in legal document formatting with specific experience in employment, 
                freelance, and NDA contracts. Skilled at creating professional, clear, 
                and legally-compliant document layouts. Specialized in creating consistent 
                formatting across different contract types.
            """),
            verbose=self.verbose,
            allow_delegation=False
        )

        # Define type-specific task templates
        contract_templates = {
            ContractType.NDA: dedent("""
                Create an NDA with:
                1. Clear definition of confidential information
                2. Scope of confidentiality obligations
                3. Duration of confidentiality period
                4. Permitted uses of information
                5. Return/destruction of confidential information
                6. Jurisdiction-specific requirements
            """),
            ContractType.FREELANCE: dedent("""
                Create a Freelance Contract with:
                1. Detailed scope of work and deliverables
                2. Payment terms and schedule
                3. Intellectual property rights
                4. Independent contractor status
                5. Term and termination conditions
                6. Liability and insurance requirements
            """),
            ContractType.EMPLOYMENT: dedent("""
                Create an Employment Contract with:
                1. Job description and duties
                2. Compensation and benefits
                3. Work schedule and location
                4. Probationary period
                5. Termination conditions
                6. Non-compete and confidentiality terms
            """)
        }

        # Define the tasks with type-specific handling
        tasks = [
            Task(
                description=dedent("""
                    Analyze the contract requirements and create initial draft based on contract type.
                    Contract Details: {text}
                    
                    Type-Specific Requirements:
                    {template}
                    
                    Additional Requirements:
                    1. Ensure all party information is correctly included
                    2. Include jurisdiction-specific clauses
                    3. Add any additional terms specified
                """),
                agent=contract_specialist,
                expected_output="Initial contract draft with type-specific components",
            ),

            Task(
                description=dedent("""
                    Review the contract draft for legal compliance based on contract type:
                    1. Verify compliance with {jurisdiction} laws
                    2. Check type-specific legal requirements
                    3. Validate all mandatory clauses
                    4. Review terminology for legal accuracy
                    
                    Previous task output: {previous_task_output}
                """),
                agent=legal_expert,
                expected_output="Legal compliance review with type-specific modifications",
            ),

            Task(
                description=dedent("""
                    Perform risk assessment for {contract_type} contract:
                    1. Identify type-specific risks and liabilities
                    2. Review jurisdiction-specific requirements
                    3. Analyze protection measures
                    4. Suggest additional safeguards
                    
                    Previous task output: {previous_task_output}
                """),
                agent=risk_analyst,
                expected_output="Risk assessment with type-specific recommendations",
            ),

            Task(
                description=dedent("""
                    Final review of {contract_type} contract:
                    1. Verify all type-specific requirements
                    2. Ensure compliance with {jurisdiction} laws
                    3. Check all risk mitigations are addressed
                    4. Validate completeness and accuracy
                    
                    Previous task output: {previous_task_output}
                """),
                agent=final_reviewer,
                expected_output="Final contract with review summary",
            ),

            Task(
                description=dedent("""
                    Format {contract_type} contract for PDF generation:
                    1. Apply type-specific formatting standards
                    2. Structure sections according to contract type
                    3. Include appropriate headers and footers
                    4. Format signature blocks
                    
                    Previous task output: {previous_task_output}
                """),
                agent=document_formatter,
                expected_output="Formatted contract ready for PDF generation",
            )
        ]

        # Create the crew with all agents and tasks
        crew = Crew(
            agents=[contract_specialist, legal_expert, risk_analyst, final_reviewer, document_formatter],
            tasks=tasks,
            verbose=self.verbose
        )

        return crew

    def process_contract(self, contract_details: ContractDetails):
        """Process contract creation and generate PDF based on contract type"""
        # Get the appropriate template
        template = self.contract_templates[contract_details.contract_type]
        
        # Prepare the context for the crew
        context = {
            "text": contract_details.dict(),
            "template": template,
            "contract_type": contract_details.contract_type,
            "jurisdiction": contract_details.jurisdiction
        }
        
        # Execute the crew tasks
        result = self.crew.kickoff(context)
        
        # Generate PDF with type-specific filename
        filename = f"{contract_details.contract_type}_{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d')}.pdf"
        pdf_path = self.generate_pdf(result, filename)
        
        return {
            "content": result,
            "pdf_path": pdf_path,
            "contract_type": contract_details.contract_type
        }
